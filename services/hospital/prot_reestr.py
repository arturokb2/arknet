from django.db import connection
from services.hospital.history import History
import docx
from hospital.models import *
from django.conf import  settings
import shutil
from asgiref.sync import async_to_sync
from channels.layers import get_channel_layer
class ProtReestrHosp:
    h = History(None)
    @classmethod
    def prot(cls,d1,d2,user,err):
        user_group_name = 'hospital_createreestr_%s' % str(user.user.id)
        file_shoblon = '/'.join([settings.MEDIA_ROOT, 'shoblons/hospital/oth', 'prot.docx'])
        file_new = '/'.join([settings.MEDIA_ROOT, 'temp', str(user.user.id), f'prot_{user.user.id}.docx'])
        shutil.copy2(file_shoblon, file_new)
        d1, d2 = cls.h.format_date(d1).replace('-', '.'), cls.h.format_date(d2).replace('-', '.')
        doc = docx.Document(file_new)
        cls._double(d1,d2,user,doc)
        doc.add_paragraph("")
        cls._err(err,doc)
        doc.save(file_new)
        async_to_sync(get_channel_layer().group_send)(user_group_name,{'type': 'download_reestr_prot', 'text': file_new})


    @staticmethod
    def _double(d1,d2,user,doc):
        cursor = connection.cursor()
        if user.statistics_type.id == 2:
            sql = f"select * from public.double_hosp('{d1}', '{d2}', '0201')"
        else:
            sql = f"select * from public.double_hosp('{d1}', '{d2}', '0202')"
        cursor.execute(sql)
        desc = cursor.description
        double = [
            dict(zip([col[0] for col in desc], row))
            for row in cursor.fetchall()
        ]
        if len(double) > 0:
            doc.add_paragraph('СПИСОК ПОВТОРНО ГОСПИТАЛИЗИРОВАННЫХ В ОТЧЕТНОМ ПЕРИОДЕ')
        for d in double:
            pat = Patient.objects.get(id=d['p_id'])
            sl = Sluchay.objects.get(id=d['s_id'])
            d1 = sl.datp.strftime('%d.%m.%Y')
            d2 = sl.datv.strftime('%d.%m.%Y')
            s = f"N ист.-{sl.nib} {pat.fam} {pat.im} {pat.ot} {sl.dskz.kod if sl.dskz else ''} {d1}-{d2} {sl.otd.naim if sl.otd else ''}"
            doc.add_paragraph(s)

    @staticmethod
    def _err(err,doc):
        err = [dict(t) for t in set([tuple(d.items()) for d in err])]
        for e in err:
            for k,v in e.items():
                o = v.sluchay.otd.naim if v.sluchay.otd else ''

                if k == 't005':
                    doc.add_paragraph(f"{v.sluchay.nib} {v.patient.fam} {v.patient.im} {v.patient.datr.strftime('%d.%m.%Y')} {o} - нет кода врача Т005")
                elif k == 'ksg_o':
                    doc.add_paragraph(f"{v.sluchay.nib} {v.patient.fam} {v.patient.im} {v.patient.datr.strftime('%d.%m.%Y')} {o} - неверный КСГ ОМС")
                elif k == 'ksg_v':
                    doc.add_paragraph(f"{v.sluchay.nib} {v.patient.fam} {v.patient.im} {v.patient.datr.strftime('%d.%m.%Y')} {o} - неверный КСГ ВЫС.")
                elif k == 'tarif':
                    doc.add_paragraph(f"{v.sluchay.nib} {v.patient.fam} {v.patient.im} {v.patient.datr.strftime('%d.%m.%Y')} {o} - пустой тариф")
                elif k == 'sumv':
                    doc.add_paragraph(f"{v.sluchay.nib} {v.patient.fam} {v.patient.im} {v.patient.datr.strftime('%d.%m.%Y')} {o} - пустая сумма")
                elif k == 't005p':
                    doc.add_paragraph(f"{v.sluchay.nib} {v.patient.fam} {v.patient.im} {v.patient.datr.strftime('%d.%m.%Y')} {o} - нет кода хир.врача Т005")
                elif k == 'vid_hmp':
                    doc.add_paragraph(f"{v.sluchay.nib} {v.patient.fam} {v.patient.im} {v.patient.datr.strftime('%d.%m.%Y')} {o} - не выставлен метод ВТМП")
                elif k == 'metod_hmp':
                    doc.add_paragraph(f"{v.sluchay.nib} {v.patient.fam} {v.patient.im} {v.patient.datr.strftime('%d.%m.%Y')} {o} - не выставлен вид ВТМП")
                elif k == 'err_text':
                    er = v.sluchay.err_text.replace('null','')
                    er = er.split(',')
                    er = [e for e in er if len(e) != 0]
                    er = set(er)
                    er = ','.join(er)
                    doc.add_paragraph(f"{v.sluchay.nib} {v.patient.fam} {v.patient.im} {v.patient.datr.strftime('%d.%m.%Y')} {o} - {er} ")




